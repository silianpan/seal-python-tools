#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2021/5/2 上午9:30
# @Author  : silianpan
# @Site    :
# @File    : file.py
# @Software: PyCharm

from openpyxl import load_workbook
from docx import Document
def info_update(doc, old_info, new_info):
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace(old_info, new_info)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_info, new_info)
wb = load_workbook('/Users/liupan/work/code/seal-python-tools/word_process/合同信息.xlsx')
ws = wb.active
for row in range(2, ws.max_row + 1):
    doc = Document('/Users/liupan/work/code/seal-python-tools/word_process/合同模板.docx')
    for col in range(1, ws.max_column + 1):
        old_info = str(ws.cell(row=1, column=col).value)
        new_info = str(ws.cell(row=row, column=col).value)
        info_update(doc, old_info, new_info)
    com_name = str(ws.cell(row=row, column=2).value)
    doc.save(f'/Users/liupan/work/code/seal-python-tools/word_process/合同文件/{com_name}合同.docx')
